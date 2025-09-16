from docx import Document
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from flask import Blueprint, abort, current_app, flash, redirect, render_template, request, send_file, url_for
from sqlalchemy import func

from ...utils.date_utils import get_period_dates
from ...utils.get_total_info import get_total_info

from ...utils.functions import combine_areas_data, combine_monthly_data, combine_monthly_data_fixed, get_monthly_data_fixed, get_monthly_data_logs, get_areas_data_applicates, get_areas_data_disinfection, get_focus_data, get_monthly_data, initialize_combined_focus_areas_data, prepare_areas_labels_and_values, prepare_data_for_chart, prepare_labels_and_values, get_top_disinfectors

from ...domain.models.user.entities import User 
from ...domain.models.application.entities import Application, Disinfection, EpidemicFocus, Area, Doctor
from ...domain.models.audit.entities import ApplicationAuditLog
from ...core.extensions import db
from flask_login import login_required, current_user
from io import BytesIO


from datetime import date, datetime, time

dashboard = Blueprint('dashboard', __name__)

@dashboard.route('/admin', methods=['POST', 'GET'])
@login_required
def admin():

    if not hasattr(current_user, 'user_type') or current_user.user_type != 'admin':
        abort(403)

    count_admins = User.query.filter_by(user_type="admin").count()
    count_doctors = User.query.filter_by(user_type="doctor").count()
    count_disinfectors = User.query.filter_by(user_type="disinfector").count()
    total_users = count_doctors + count_disinfectors + count_admins
    total_applications = Application.query.count()
    total_changes = ApplicationAuditLog.query.count()

    year = datetime.now().year
    period = request.args.get('period', default='год')

    start_date, end_date = get_period_dates(period, year)

    monthly_data_changes = get_monthly_data_logs(
        ApplicationAuditLog, 'change_time', start_date, end_date)

    audit_logs = (
        db.session.query(
            ApplicationAuditLog,
            User.login.label("login"),
            User.user_type.label("user_type"),
        )
        .outerjoin(User, User.id == ApplicationAuditLog.changed_by)
        .order_by(ApplicationAuditLog.change_time.desc())
        .all()
    )
    
    
    months = list(monthly_data_changes.keys())
    insert_data = [monthly_data_changes[month]['INSERT'] for month in months]
    update_data = [monthly_data_changes[month]['UPDATE'] for month in months]
    delete_data = [monthly_data_changes[month]['DELETE'] for month in months]

    total_info = {
        'total_users': total_users,
        'count_admins': count_admins,
        'count_doctors': count_doctors,
        'count_disinfectors': count_disinfectors,
        'total_applications': total_applications,
        'total_changes': total_changes
    }
    
    query_apllications = (db.session.query(Application, Disinfection, Area, Doctor).join(Disinfection, Disinfection.application_id == Application.id).join(Doctor, Doctor.id == Application.user_id).join(Area, Area.id == Doctor.area_id))

    applications = query_apllications.order_by(Application.submission_date.asc()).all()

    return render_template(
        'application/admin.html',
        total_info=total_info,
        monthly_data_changes=monthly_data_changes,
        months=months,
        insert_data=insert_data,
        update_data=update_data,
        delete_data=delete_data,
        period=period,
        audit_logs=audit_logs,
        applications = applications
    )



@dashboard.route('/dashboard', methods=['POST', 'GET'])
@login_required
def all():

    year = datetime.now().year
    
    period = request.args.get('period', default='год')

    start_date, end_date = get_period_dates(period, year)
    
    total_info = get_total_info(start_date, end_date)

    # monthly_data_applicates = get_monthly_data(Application, 'submission_date', start_date, end_date)
    # monthly_data_disinfection = get_monthly_data(Disinfection, 'disinfection_date', start_date, end_date)

    # combined_data = combine_monthly_data(monthly_data_applicates, monthly_data_disinfection, start_date, end_date)
    
    monthly_data_applicates = get_monthly_data_fixed(Application, 'submission_date', start_date, end_date)
    monthly_data_disinfection = get_monthly_data_fixed(Disinfection, 'disinfection_date', start_date, end_date)
    combined_data = combine_monthly_data_fixed(monthly_data_applicates, monthly_data_disinfection, start_date, end_date)



    all_months = list(combined_data.keys())
    total_values_full = [data['applications'] for data in combined_data.values()]
    completed_values_full = [data['disinfections'] for data in combined_data.values()]

    areas_data_applicates = get_areas_data_applicates(start_date, end_date)
    areas_data_disinfection = get_areas_data_disinfection(start_date, end_date)
    combined_areas_data = combine_areas_data(areas_data_applicates, areas_data_disinfection)
    areas_labels, areas_total_values, areas_completed_values = prepare_areas_labels_and_values(combined_areas_data)

    focus_data = get_focus_data(start_date, end_date)
    combined_focus_areas_data = initialize_combined_focus_areas_data(focus_data)
    focus_labels, area_labelsf, focus_total_values = prepare_labels_and_values(combined_focus_areas_data)
    focus_total_values_list = prepare_data_for_chart(focus_total_values, focus_labels)
    
    all_areas = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', 'ДПДО', 'РПТД']
    epidemic_groups = ['I', 'II', 'III', 'IV', 'V', 'безадресные']

    query = db.session.query(EpidemicFocus.name.label('epidemic_focus'),Area.name_area.label('area'),func.count(Application.id).label('count')).join(Application, Application.focus_id == EpidemicFocus.id).outerjoin(Doctor, Doctor.id == Application.user_id).outerjoin(Area, Area.id == Doctor.area_id).filter(Application.submission_date.between(start_date, end_date)).filter(Area.name_area.in_(all_areas)).group_by(EpidemicFocus.name, Area.name_area)
    
    results = query.all()

    posthumous_query = db.session.query(EpidemicFocus.name.label('epidemic_focus'),Area.name_area.label('area'),func.count(Application.id).label('posthumous_count')).join(Application, Application.focus_id == EpidemicFocus.id).outerjoin(Doctor, Doctor.id == Application.user_id).outerjoin(Area, Area.id == Doctor.area_id).filter(Application.submission_date.between(start_date, end_date)).filter(Application.reason_application == 'posthumously').filter(Area.name_area.in_(all_areas)).group_by(EpidemicFocus.name, Area.name_area)
    
    posthumous_results = posthumous_query.all()

    table_data = []
    col_totals = {area: 0 for area in all_areas}
    posthumous_col_totals = {area: 0 for area in all_areas}

    for group in epidemic_groups:
        row = {'focus': group}
        row_total = 0
        
        for area in all_areas:
            count = next((r.count for r in results 
                         if r.epidemic_focus == group and r.area == area), 0)
            row[area] = count
            row_total += count
            col_totals[area] += count
            
            posthumous_count = next((r.posthumous_count for r in posthumous_results 
                                   if r.epidemic_focus == group and r.area == area), 0)
            posthumous_col_totals[area] += posthumous_count
        
        row['total'] = row_total
        table_data.append(row)
    
    posthumous_row = {'focus': 'в том числе посмертно'}
    for area in all_areas:
        posthumous_row[area] = posthumous_col_totals[area]
    posthumous_row['total'] = sum(posthumous_col_totals.values())
    table_data.append(posthumous_row)
    
    total_row = {'focus': 'ВСЕГО'}
    for area in all_areas:
        total_row[area] = col_totals[area]
    total_row['total'] = sum(col_totals.values())
    table_data.append(total_row)
    
    
    top_disinfectors = get_top_disinfectors(start_date, end_date)
    
    return render_template(
        'application/dashboard.html',
        total_info=total_info,
        months=all_months, total_values=total_values_full, completed_values=completed_values_full,
        areas_labels=areas_labels,
        areas_total_values=areas_total_values,
        areas_completed_values=areas_completed_values,
        focus_labels=focus_labels,
        area_labelsf=area_labelsf,
        focus_total_values_list=focus_total_values_list,
        period=period, table_data=table_data, areas=all_areas, top_disinfectors=top_disinfectors)


@dashboard.route('/generate_word_report')
def generate_word_report():
    try:
        year = datetime.now().year
        period = request.args.get('period', default='год')
        
        start_date, end_date = get_period_dates(period, year)
        end_date_with_time = datetime.combine(end_date, time.max) if isinstance(end_date, date) else end_date
        
        all_areas = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', 'ДПДО', 'РПТД']
        epidemic_groups = ['I', 'II', 'III', 'IV', 'V', 'безадресные']

        query = (
            db.session.query(
                EpidemicFocus.name.label("epidemic_focus"),
                Area.name_area.label("area"),
                func.count(Application.id).label("count"),
            )
            .join(Application, Application.focus_id == EpidemicFocus.id)
            .outerjoin(Doctor, Doctor.id == Application.doctor_id)
            .outerjoin(Area, Area.id == Doctor.area_id)
            .filter(Application.submission_date.between(start_date, end_date))
            .filter(Area.name_area.in_(all_areas))
            .group_by(EpidemicFocus.name, Area.name_area)
        )
        
        results = query.all()
        
        posthumous_query = (
            db.session.query(
                EpidemicFocus.name.label("epidemic_focus"),
                Area.name_area.label("area"),
                func.count(Application.id).label("posthumous_count"),
            )
            .join(Application, Application.focus_id == EpidemicFocus.id)
            .outerjoin(Doctor, Doctor.id == Application.doctor_id)
            .outerjoin(Area, Area.id == Doctor.area_id)
            .filter(Application.submission_date.between(start_date, end_date))
            .filter(Application.reason_application == "posthumously")
            .filter(Area.name_area.in_(all_areas))
            .group_by(EpidemicFocus.name, Area.name_area)
        )
        
        posthumous_results = posthumous_query.all()

        table_data = []
        col_totals = {area: 0 for area in all_areas}
        posthumous_col_totals = {area: 0 for area in all_areas}
        
        for group in epidemic_groups:
            row = {'focus': group}
            row_total = 0
            
            for area in all_areas:
                count = next((r.count for r in results 
                            if r.epidemic_focus == group and r.area == area), 0)
                row[area] = count
                row_total += count
                col_totals[area] += count
                
                posthumous_count = next((r.posthumous_count for r in posthumous_results 
                                      if r.epidemic_focus == group and r.area == area), 0)
                posthumous_col_totals[area] += posthumous_count
            
            row['total'] = row_total
            table_data.append(row)
        
        posthumous_row = {'focus': 'в том числе посмертно'}
        for area in all_areas:
            posthumous_row[area] = posthumous_col_totals[area]
        posthumous_row['total'] = sum(posthumous_col_totals.values())
        table_data.append(posthumous_row)
        
        total_row = {'focus': 'ВСЕГО'}
        for area in all_areas:
            total_row[area] = col_totals[area]
        total_row['total'] = sum(col_totals.values())
        table_data.append(total_row)
        
        total_info = get_total_info(start_date, end_date_with_time)
        count_posthumously = Application.query.filter(
            Application.submission_date >= start_date,
            Application.submission_date <= end_date_with_time,
            Application.reason_application == 'posthumously'
        ).count()
        
        count_refusal = Application.query.filter(
            Application.submission_date >= start_date,
            Application.submission_date <= end_date_with_time,
            Application.status == 'refusal'
        ).count()
        
        square_disinfection = db.session.query(func.sum(Disinfection.area_size))\
            .filter(
                Disinfection.disinfection_date >= start_date,
                Disinfection.disinfection_date <= end_date_with_time
            ).scalar() or 0
        
        doc = Document()
        
        title = doc.add_heading('Отчет по заключительной дезинфекции в очагах туберкулеза', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        period_paragraph = doc.add_paragraph()
        period_paragraph.add_run(f"Период: {start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}").bold = True
        
        doc.add_paragraph() 
        
        doc.add_heading("Исполнение заявок на заключительную дезинфекцию", level=2)
        
        info_items = [
            f"Поступило заявок всего: {total_info['total']}",
            f"в том числе: посмертно: {count_posthumously}",
            f"Отказ: {count_refusal}",
            f"Обработано кв.м.: {square_disinfection}",
            f"Взято смывов: {total_info.get('washes', 0)}"
        ]
        
        for item in info_items:
            doc.add_paragraph(item)
        
        doc.add_paragraph(f"ИТОГО: {total_info['completed']}", style='Heading3')
        doc.add_paragraph()
        
        doc.add_heading("Количество заявок на заключительную дезинфекцию в разрезе противотуберкулезных диспансеров", level=2)
        table = doc.add_table(rows=1, cols=len(all_areas)+2)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ['Группы очагов/участки'] + all_areas + ['ВСЕГО']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for row in table_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row['focus']
            row_cells[0].paragraphs[0].runs[0].font.bold = row['focus'] in ['в том числе посмертно', 'ВСЕГО']
            
            for i, area in enumerate(all_areas, 1):
                row_cells[i].text = str(row.get(area, 0))
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_cells[-1].text = str(row.get('total', 0))
            row_cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        date_paragraph = doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        byte_stream = BytesIO()
        doc.save(byte_stream)
        byte_stream.seek(0)
        
        filename = f"epid_report_{start_date.strftime('%Y-%m-%d')}_to_{end_date.strftime('%Y-%m-%d')}.docx"
        return send_file(
            byte_stream,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        current_app.logger.error(f"Error generating report: {str(e)}", exc_info=True)
        flash('Произошла ошибка при генерации отчета', 'danger')
        return redirect(url_for('dashboard.all'))
